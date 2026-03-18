from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple

from docx import Document

# ---------- Helpers ----------
def norm(s: str) -> str:
    return re.sub(r"\s+", " ", (s or "").replace("\xa0", " ")).strip()


def is_blank(s: str) -> bool:
    return not norm(s)


# Nhận diện bắt đầu câu hỏi:
# - "Câu 1:" / "Câu 1." / "Câu 1 -"
# - "1." / "1)" / "1 -"
RE_Q_PREFIX = re.compile(r"^\s*(?:Câu\s*)?(\d+)\s*([.)\:\-–])\s*", re.IGNORECASE)

# Nhận diện option dạng dòng riêng
RE_OPT_LINE = re.compile(r"^\s*([A-Da-d])\s*([.)\:\-–])\s*(.+?)\s*$")

# Tách option trong 1 dòng inline: tìm các mốc "A." "B." "C." "D."
RE_OPT_MARK = re.compile(r"\b([A-D])\s*([.)\:\-–])\s*", re.IGNORECASE)


@dataclass
class RunSpan:
    start: int
    end: int
    bold: bool


def paragraph_spans(paragraph) -> Tuple[str, List[RunSpan]]:
    """
    Build a plain text string + spans mapping run index ranges into that text,
    preserving concatenation order, so we can locate bold positions.
    """
    parts: List[str] = []
    spans: List[RunSpan] = []
    pos = 0
    for run in paragraph.runs:
        t = run.text or ""
        if not t:
            continue
        parts.append(t)
        start = pos
        pos += len(t)
        spans.append(RunSpan(start=start, end=pos, bold=bool(run.bold)))
    text = "".join(parts)
    return text, spans


def any_bold_in_range(spans: List[RunSpan], start: int, end: int) -> bool:
    for sp in spans:
        if sp.end <= start:
            continue
        if sp.start >= end:
            break
        if sp.bold:
            return True
    return False


# ---------- Image extraction ----------
# map common Office image content-types to file extension
CT_TO_EXT = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/jpg": "jpg",
    "image/gif": "gif",
    "image/bmp": "bmp",
    "image/tiff": "tiff",
    "image/webp": "webp",
    "image/x-emf": "emf",
    "image/emf": "emf",
    "image/x-wmf": "wmf",
    "image/wmf": "wmf",
}


def _content_type_to_ext(content_type: Optional[str]) -> str:
    if not content_type:
        return "bin"
    return CT_TO_EXT.get(content_type.lower(), "bin")


def extract_question_images(
    doc: Document,
    subject_code: str,
    out_dir: str,
) -> Dict[int, List[str]]:
    """
    Stable mapping by paragraph windows:
      - Determine question start paragraph indices by RE_Q_PREFIX on doc.paragraphs
      - For each paragraph in [start_idx_of_q .. start_idx_of_next_q), collect all r:embed rIds
      - Export:
          * If 1 image: <subject>q<qno>.<ext>
          * If >=2 images: <subject>q<qno>(n).<ext> with n=1..k
    """
    os.makedirs(out_dir, exist_ok=True)

    # 1) find question start paragraph indices
    q_starts: List[Tuple[int, int]] = []  # (qno, start_para_idx)
    for idx, p in enumerate(doc.paragraphs):
        t = norm(p.text)
        if not t:
            continue
        m = RE_Q_PREFIX.match(t)
        if m:
            qno = int(m.group(1))
            q_starts.append((qno, idx))

    if not q_starts:
        return {}

    q_starts.sort(key=lambda x: x[1])

    # 2) build paragraph xml list once
    para_xml: List[str] = [p._p.xml for p in doc.paragraphs]  # noqa: SLF001

    # 3) collect rIds in each question window
    qno_to_rids: Dict[int, List[str]] = {}
    for i, (qno, start_idx) in enumerate(q_starts):
        end_idx = q_starts[i + 1][1] if i + 1 < len(q_starts) else len(doc.paragraphs)

        rids: List[str] = []
        seen = set()
        for j in range(start_idx, end_idx):
            xml = para_xml[j]
            for rid in re.findall(r'r:embed="(rId\d+)"', xml):
                if rid not in seen:
                    seen.add(rid)
                    rids.append(rid)

        if rids:
            qno_to_rids[qno] = rids

    # 4) export files with your naming rule
    qno_to_paths: Dict[int, List[str]] = {}
    for qno, rids in qno_to_rids.items():
        total = len(rids)
        for idx, rid in enumerate(rids, start=1):
            part = doc.part.related_parts.get(rid)
            if not part:
                continue

            ext = _content_type_to_ext(getattr(part, "content_type", None))

            if total == 1:
                filename = f"{subject_code}q{qno}.{ext}"
            else:
                filename = f"{subject_code}q{qno}({idx}).{ext}"

            save_path = os.path.join(out_dir, filename)
            with open(save_path, "wb") as f:
                f.write(part.blob)

            rel = os.path.join("static", "pic", filename).replace("/", "/")
            qno_to_paths.setdefault(qno, []).append(rel)

    return qno_to_paths

def attach_images(item: Dict, qno: Optional[int], qno_to_imgs: Optional[Dict[int, List[str]]]):
    if not qno_to_imgs or not qno:
        return
    imgs = qno_to_imgs.get(qno) or []
    if not imgs:
        return
    item["imgs"] = imgs
    # Backward compatibility: keep first image as "img"
    item["img"] = imgs[0]


# ---------- Mode: default (multi-line) ----------
def parse_default(doc: Document, qno_to_imgs: Optional[Dict[int, List[str]]] = None) -> List[Dict]:
    quiz: List[Dict] = []
    current: Optional[Dict] = None
    in_options = False
    current_qno: Optional[int] = None

    def finalize():
        nonlocal current, in_options, current_qno
        if not current:
            return
        # ensure keys
        current.setdefault("opts", {})
        for k in ["A", "B", "C", "D"]:
            current["opts"].setdefault(k, "")
        opt_count = sum(1 for k in ["A", "B", "C", "D"] if norm(current["opts"][k]))
        if norm(current.get("q", "")) and opt_count >= 2:
            attach_images(current, current_qno, qno_to_imgs)
            quiz.append(current)
        current = None
        in_options = False
        current_qno = None

    for p in doc.paragraphs:
        line = norm(p.text)
        if not line:
            continue

        # start new question
        m = RE_Q_PREFIX.match(line)
        if m:
            finalize()
            current_qno = int(m.group(1))
            # remove prefix
            q_text = norm(line[m.end():])
            current = {
                "id": 0,
                "q": q_text or line,
                "opts": {"A": "", "B": "", "C": "", "D": ""},
                "ans": "",
                "type": "mcq",
            }
            in_options = False
            continue

        if current is None:
            # fallback: treat as question (no number)
            current = {
                "id": 0,
                "q": line,
                "opts": {"A": "", "B": "", "C": "", "D": ""},
                "ans": "",
                "type": "mcq",
            }
            in_options = False
            current_qno = None
            continue

        # option line
        om = RE_OPT_LINE.match(line)
        if om:
            in_options = True
            key = om.group(1).upper()
            text = norm(om.group(3))
            current["opts"][key] = text

            # answer: if any bold run in this paragraph -> that option correct
            if any(r.bold and norm(r.text) for r in p.runs):
                current["ans"] = key
            continue

        # continuation line
        if not in_options:
            current["q"] = norm(current["q"] + " " + line)
        else:
            # append to last option if exists
            last_key = None
            for k in ["D", "C", "B", "A"]:
                if norm(current["opts"][k]):
                    last_key = k
                    break
            if last_key:
                current["opts"][last_key] = norm(current["opts"][last_key] + " " + line)
                if any(r.bold and norm(r.text) for r in p.runs) and not current.get("ans"):
                    current["ans"] = last_key
            else:
                current["q"] = norm(current["q"] + " " + line)

    finalize()

    for i, item in enumerate(quiz, start=1):
        item["id"] = i
    return quiz


# ---------- Mode: inline (single paragraph has q + options) ----------
def split_inline_question_and_options(
    text: str,
) -> Optional[Tuple[str, Dict[str, str], Dict[str, Tuple[int, int]]]]:
    matches = list(RE_OPT_MARK.finditer(text))
    if len(matches) < 2:
        return None

    segments = []
    for i, m in enumerate(matches):
        key = m.group(1).upper()
        start = m.start()
        content_start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        segments.append((key, start, content_start, end))

    keys = [k for (k, *_rest) in segments]
    if "A" not in keys or "B" not in keys:
        return None

    q_text = norm(text[:segments[0][1]])
    opts: Dict[str, str] = {"A": "", "B": "", "C": "", "D": ""}
    ranges: Dict[str, Tuple[int, int]] = {}

    for key, start, content_start, end in segments:
        opt_text = norm(text[content_start:end])
        if key in opts:
            opts[key] = opt_text
            ranges[key] = (start, end)

    opt_count = sum(1 for k in ["A", "B", "C", "D"] if norm(opts[k]))
    if not q_text or opt_count < 2:
        return None

    return q_text, opts, ranges


def parse_inline(doc: Document, qno_to_imgs: Optional[Dict[int, List[str]]] = None) -> List[Dict]:
    quiz: List[Dict] = []

    for p in doc.paragraphs:
        raw_text, spans = paragraph_spans(p)
        text = norm(raw_text)
        if not text:
            continue

        qno = None
        m_qno = RE_Q_PREFIX.match(text)
        if m_qno:
            qno = int(m_qno.group(1))

        m = RE_Q_PREFIX.match(text)
        if m:
            raw = raw_text
            m2 = RE_Q_PREFIX.match(raw)
            if m2:
                raw_wo_prefix = raw[m2.end():]
                split = split_inline_question_and_options(raw_wo_prefix)
                if not split:
                    split = split_inline_question_and_options(raw)
                    if not split:
                        continue
                    q_text, opts, opt_ranges = split
                    prefix_offset = 0
                else:
                    q_text, opts, opt_ranges = split
                    prefix_offset = m2.end()
            else:
                split = split_inline_question_and_options(raw_text)
                if not split:
                    continue
                q_text, opts, opt_ranges = split
                prefix_offset = 0
        else:
            raw = raw_text
            split = split_inline_question_and_options(raw)
            if not split:
                continue
            q_text, opts, opt_ranges = split
            prefix_offset = 0

        ans = ""
        for key in ["A", "B", "C", "D"]:
            if key not in opt_ranges:
                continue
            start, end = opt_ranges[key]
            start += prefix_offset
            end += prefix_offset
            if any_bold_in_range(spans, start, end):
                ans = key
                break

        item = {
            "id": 0,
            "q": norm(q_text),
            "opts": {k: norm(v) for k, v in opts.items()},
            "ans": ans,
            "type": "mcq",
        }

        attach_images(item, qno, qno_to_imgs)
        quiz.append(item)

    out = []
    for i, item in enumerate(quiz, start=1):
        item["id"] = i
        out.append(item)
    return out


# ---------- Auto mode ----------
def detect_mode(doc: Document) -> str:
    inline_hits = 0
    optline_hits = 0
    checked = 0

    for p in doc.paragraphs:
        t = norm(p.text)
        if not t:
            continue
        checked += 1
        if split_inline_question_and_options(t):
            inline_hits += 1
        if RE_OPT_LINE.match(t):
            optline_hits += 1
        if checked >= 50:
            break

    if inline_hits >= 2 and inline_hits >= optline_hits:
        return "inline"
    return "default"


def docx_to_quiz(
    docx_path: str,
    mode: str,
    subject_code: str = "",
    image_dir: str = os.path.join("static", "pic"),
) -> List[Dict]:
    doc = Document(docx_path)

    qno_to_imgs: Dict[int, List[str]] = {}
    if subject_code:
        qno_to_imgs = extract_question_images(doc, subject_code=subject_code, out_dir=image_dir)

    if mode == "auto":
        mode = detect_mode(doc)

    if mode == "default":
        return parse_default(doc, qno_to_imgs=qno_to_imgs)
    if mode == "inline":
        return parse_inline(doc, qno_to_imgs=qno_to_imgs)

    raise ValueError("mode must be one of: default, inline, auto")


# ---------- CLI ----------
def main():
    ap = argparse.ArgumentParser(description="Scan DOCX -> JSON quiz (2 mode).")
    ap.add_argument("input", help="Đường dẫn file .docx")
    ap.add_argument("--mode", choices=["default", "inline", "auto"], default="auto")
    ap.add_argument("-o", "--output", default="", help="File output .json (mặc định: stdout)")
    ap.add_argument("--pretty", action="store_true", help="In JSON đẹp (indent=2, UTF-8)")

    ap.add_argument(
        "--subject",
        default="",
        help="Mã môn để đặt tên ảnh: <subject>q<so_cau>(n).* (bỏ trống = không export ảnh)",
    )
    ap.add_argument(
        "--image-dir",
        default=os.path.join("static", "pic"),
        help=r"Thư mục lưu ảnh (mặc định: static\pic\)",
    )

    args = ap.parse_args()

    quiz = docx_to_quiz(
        args.input,
        args.mode,
        subject_code=args.subject,
        image_dir=args.image_dir,
    )

    s = json.dumps(quiz, ensure_ascii=False, indent=2 if args.pretty else None)
    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(s)
    else:
        print(s)


if __name__ == "__main__":
    main()
    #cách chạy: python back\tools.py input.docx --mode auto --subject xx123 -o quiz.json --pretty